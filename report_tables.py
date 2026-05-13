"""Helpers for populating DOCX template tables in clinical reports.

Split out from generate_clinical_reports.py to keep that module under
the 300-line file-size limit. These helpers operate on python-docx
Table objects directly and contain only presentation/formatting logic.
"""

from __future__ import annotations

from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# @AX:NOTE [AUTO] 12-tuple mapping is the runtime mirror of spec.md Table 0 mapping;
# @AX:NOTE [AUTO] any edit here must be reflected in spec.md and acceptance S5.
# Table 0: 12 cells (4 rows x 3 label/value pairs). Mapping per SPEC-REPORT-001.
# Each tuple is (row_index, cell_index, meta_key).
TABLE0_MAPPING = [
    (0, 1, "patient_name"),
    (0, 3, "specimen_type"),
    (0, 5, "test_no"),
    (1, 1, "reg_no"),
    (1, 3, "specimen_collected_at"),
    (1, 5, "test_date"),
    (2, 1, "birth_date"),
    (2, 3, "specimen_received_at"),
    (2, 5, "preliminary_report_date"),
    (3, 1, "ordering_doctor"),
    (3, 3, "specimen_state"),
    (3, 5, "final_report_date"),
]


def set_cell_text(cell, text, font_size=Pt(9), bold=False,
                  align=WD_ALIGN_PARAGRAPH.CENTER):
    """Set cell text with formatting, clearing existing content."""
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.size = font_size
    run.font.bold = bold
    run.font.name = "Malgun Gothic"


def fill_table0_patient_info(table, meta: dict) -> None:
    """Populate Table 0's 12 patient-info cells from a meta.json dict.

    Missing keys yield empty cells (REQ-6 fail-soft).
    """
    sz = Pt(9)
    for row_idx, col_idx, key in TABLE0_MAPPING:
        value = meta.get(key, "")
        set_cell_text(table.rows[row_idx].cells[col_idx], value, sz)


def fill_table14_qc_data(table, qc: dict) -> None:
    """Populate Table 14 (3 rows x 2 cols) with QC depth statistics.

    Expects qc dict with keys: mean_depth, median_depth, pct_dp_ge_100x.
    """
    qc_pairs = [
        ("Mean Depth",   f"{qc['mean_depth']:.1f}x"),
        ("Median Depth", f"{qc['median_depth']:.1f}x"),
        ("% DP ≥ 100x", f"{qc['pct_dp_ge_100x']:.1f}%"),
    ]
    for ri, (label, value) in enumerate(qc_pairs):
        set_cell_text(
            table.rows[ri].cells[0], label, Pt(9),
            bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
        )
        set_cell_text(
            table.rows[ri].cells[1], value, Pt(9),
            bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
        )


def fill_table15_signatures(table, meta: dict) -> None:
    """Populate Table 15 (1 row x 5 cells) with examiner / reporter names.

    Format follows test_report_table15.py expectations:
      cell[0] = '검사자'
      cell[1] = examiners joined with '/'
      cell[2] = '' (template middle gap)
      cell[3] = '보고자'
      cell[4] = reporters joined with '/'
    """
    examiners = meta.get("examiners") or ["", ""]
    reporters = meta.get("reporters") or ["", "", "", ""]
    row0 = table.rows[0]
    set_cell_text(row0.cells[0], "검사자", Pt(9))
    set_cell_text(
        row0.cells[1],
        "/".join(str(x) for x in examiners[:2]),
        Pt(9),
    )
    set_cell_text(row0.cells[2], "", Pt(9))
    set_cell_text(row0.cells[3], "보고자", Pt(9))
    set_cell_text(
        row0.cells[4],
        "/".join(str(x) for x in reporters[:4]),
        Pt(9),
    )


def vaf_float(row: dict) -> float:
    """Parse '% VAF' field (e.g. '12.5%') as a float; 0.0 on failure."""
    try:
        return float(row.get("% VAF", "0%").rstrip("%"))
    except (ValueError, AttributeError):
        return 0.0
