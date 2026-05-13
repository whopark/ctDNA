"""
Phase 1.5 test scaffold — S5: meta.json → Table 0 six slots + Table 15 populated.

Integration test that:
1. Creates fixture meta.json, tiered CSV, annotated CSV in a tmp directory
2. Calls generate_report() to produce a DOCX
3. Opens the DOCX and asserts Table 0 and Table 15 cell contents

RED state: generate_report currently does NOT read meta.json for Table 0/14/15,
so all Table 0 and Table 15 assertions will fail at the assertion level (not ImportError).
"""
from __future__ import annotations

import csv
import json
import os
import sys
from pathlib import Path

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, ROOT)

from generate_clinical_reports import generate_report  # noqa: E402


# ---- Fixture helpers --------------------------------------------------------

META_FIXTURE = {
    "patient_name": "홍길동",
    "reg_no": "2026-TEST-01",
    "test_no": "TEST-001",
    "specimen_type": "Peripheral Blood",
    "specimen_state": "Fresh",
    "test_date": "2026-03-25",
    "birth_date": "1968-04-12",
    "ordering_doctor": "김의사",
    "specimen_collected_at": "2026-03-20",
    "specimen_received_at": "2026-03-22",
    "preliminary_report_date": "2026-03-26",
    "final_report_date": "2026-03-28",
    "examiners": ["정민경", "하명오"],
    "reporters": ["조성현", "최용준", "박주헌", "신명근"],
}


def _write_tiered_csv(path: Path) -> str:
    """Write a minimal tiered report CSV with one Tier 1 variant."""
    fieldnames = [
        "ASCO/AMP Classification", "Gene", "Canonical name",
        "Accession", "Nucleotide change", "AA change", "% VAF",
    ]
    with open(str(path), "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerow({
            "ASCO/AMP Classification": "Tier 1",
            "Gene": "MYD88",
            "Canonical name": "MYD88",
            "Accession": "NM_002468",
            "Nucleotide change": "c.818T>C",
            "AA change": "p.Leu273Pro",
            "% VAF": "12.5%",
        })
    return str(path)


def _write_annotated_csv(path: Path) -> str:
    """Write a minimal annotated CSV."""
    fieldnames = ["chrom", "pos", "ref", "alt", "dp", "tier"]
    with open(str(path), "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerow({
            "chrom": "3", "pos": "38182641", "ref": "T", "alt": "C",
            "dp": "200", "tier": "Tier 1",
        })
    return str(path)


def _get_cell_text(cell) -> str:
    """Extract all text from a python-docx cell."""
    return "".join(p.text for p in cell.paragraphs).strip()


# ---- Test class -------------------------------------------------------------

class TestReportTable15S5:
    """S5: when meta.json present, Table 0 and Table 15 are fully populated."""

    def _setup_case(self, tmp_path: Path) -> tuple[str, str, str, str]:
        """Create all fixture files and return (case_id, tiered_csv, annotated_csv, out_path)."""
        case_id = "01-TEST_99999"
        case_dir = tmp_path / case_id
        case_dir.mkdir()

        # Write meta.json
        with open(str(case_dir / "meta.json"), "w", encoding="utf-8") as f:
            json.dump(META_FIXTURE, f, ensure_ascii=False)

        tiered_csv = _write_tiered_csv(case_dir / f"{case_id}_tiered_report.csv")
        annotated_csv = _write_annotated_csv(case_dir / f"{case_id}_annotated.csv")
        output_path = str(tmp_path / f"{case_id}_clinical_report.docx")
        return case_id, tiered_csv, annotated_csv, output_path

    def test_table0_birth_date_not_empty(self, tmp_path):
        """Table 0 R2c1 must be non-empty when meta.json has birth_date."""
        from docx import Document
        case_id, tiered_csv, annotated_csv, output_path = self._setup_case(tmp_path)
        generate_report(case_id, tiered_csv, annotated_csv, output_path)
        assert os.path.exists(output_path), "generate_report must produce a .docx file"
        doc = Document(output_path)
        t0 = doc.tables[0]
        cell_text = _get_cell_text(t0.rows[2].cells[1])
        assert cell_text != "", (
            "Table 0 R2c1 (birth_date) must not be empty when meta.json is present"
        )

    def test_table15_examiner_cell(self, tmp_path):
        """Table 15 row 0: cell[1] must read '정민경/하명오'."""
        from docx import Document
        case_id, tiered_csv, annotated_csv, output_path = self._setup_case(tmp_path)
        generate_report(case_id, tiered_csv, annotated_csv, output_path)
        doc = Document(output_path)
        t15 = doc.tables[15]
        row0 = t15.rows[0]
        cell0_text = _get_cell_text(row0.cells[0])
        cell1_text = _get_cell_text(row0.cells[1])
        assert cell0_text == "검사자", (
            f"Table 15 row0 cell[0] must be '검사자', got {cell0_text!r}"
        )
        assert cell1_text == "정민경/하명오", (
            f"Table 15 row0 cell[1] must be '정민경/하명오', got {cell1_text!r}"
        )

    def test_table15_reporter_cell(self, tmp_path):
        """Table 15 row 0: cell[3]='보고자', cell[4]='조성현/최용준/박주헌/신명근'."""
        from docx import Document
        case_id, tiered_csv, annotated_csv, output_path = self._setup_case(tmp_path)
        generate_report(case_id, tiered_csv, annotated_csv, output_path)
        doc = Document(output_path)
        t15 = doc.tables[15]
        row0 = t15.rows[0]
        cell3_text = _get_cell_text(row0.cells[3])
        cell4_text = _get_cell_text(row0.cells[4])
        assert cell3_text == "보고자", (
            f"Table 15 row0 cell[3] must be '보고자', got {cell3_text!r}"
        )
        assert cell4_text == "조성현/최용준/박주헌/신명근", (
            f"Table 15 row0 cell[4] must be '조성현/최용준/박주헌/신명근', got {cell4_text!r}"
        )

    def test_table0_birth_date(self, tmp_path):
        """Table 0 row 2 cell 1 must read '1968-04-12' from meta.json."""
        from docx import Document
        case_id, tiered_csv, annotated_csv, output_path = self._setup_case(tmp_path)
        generate_report(case_id, tiered_csv, annotated_csv, output_path)
        doc = Document(output_path)
        t0 = doc.tables[0]
        cell_text = _get_cell_text(t0.rows[2].cells[1])
        assert cell_text == "1968-04-12", (
            f"Table 0 R2c1 must be '1968-04-12' (birth_date), got {cell_text!r}"
        )

    def test_table0_ordering_doctor(self, tmp_path):
        """Table 0 row 3 cell 1 must read '김의사' from meta.json."""
        from docx import Document
        case_id, tiered_csv, annotated_csv, output_path = self._setup_case(tmp_path)
        generate_report(case_id, tiered_csv, annotated_csv, output_path)
        doc = Document(output_path)
        t0 = doc.tables[0]
        cell_text = _get_cell_text(t0.rows[3].cells[1])
        assert cell_text == "김의사", (
            f"Table 0 R3c1 must be '김의사' (ordering_doctor), got {cell_text!r}"
        )

    def test_table0_preliminary_report_date(self, tmp_path):
        """Table 0 row 2 cell 5 must read preliminary_report_date, not test_date."""
        from docx import Document
        case_id, tiered_csv, annotated_csv, output_path = self._setup_case(tmp_path)
        generate_report(case_id, tiered_csv, annotated_csv, output_path)
        doc = Document(output_path)
        t0 = doc.tables[0]
        cell_text = _get_cell_text(t0.rows[2].cells[5])
        assert cell_text == "2026-03-26", (
            f"Table 0 R2c5 must be '2026-03-26' (preliminary_report_date), "
            f"not test_date; got {cell_text!r}"
        )

    def test_table0_final_report_date(self, tmp_path):
        """Table 0 row 3 cell 5 must read final_report_date, not test_date."""
        from docx import Document
        case_id, tiered_csv, annotated_csv, output_path = self._setup_case(tmp_path)
        generate_report(case_id, tiered_csv, annotated_csv, output_path)
        doc = Document(output_path)
        t0 = doc.tables[0]
        cell_text = _get_cell_text(t0.rows[3].cells[5])
        assert cell_text == "2026-03-28", (
            f"Table 0 R3c5 must be '2026-03-28' (final_report_date), "
            f"not test_date; got {cell_text!r}"
        )
