#!/usr/bin/env python3
"""
Generate clinical .docx reports from tiered CSV + annotated CSV data,
using the KBB report_template.docx as the base template.
"""

import csv
import copy
import sys
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

TEMPLATE = os.path.join(os.path.dirname(__file__),
    '..', '.agents', 'skills', 'kbb-annotation', 'reference', 'report_template.docx')

# Case metadata
CASES = {
    "01-JJH_10679562": {
        "patient": "JJH",
        "reg_no": "10679562",
        "specimen": "Peripheral Blood",
        "test_date": "2026-02-24",
    },
    "02-OKW_10673102": {
        "patient": "OKW",
        "reg_no": "10673102",
        "specimen": "Peripheral Blood",
        "test_date": "2026-02-24",
    },
    "03-OJS_23953884": {
        "patient": "OJS",
        "reg_no": "23953884",
        "specimen": "Peripheral Blood",
        "test_date": "2026-02-24",
    },
    "04-PHJ_10680696": {
        "patient": "PHJ",
        "reg_no": "10680696",
        "specimen": "Peripheral Blood",
        "test_date": "2026-02-24",
    },
}

# Therapeutic implications lookup
THERAPEUTIC = {
    "EZH2": "Tazemetostat (FDA-approved, FL; R/R DLBCL data)",
    "PTEN": "PI3K/AKT/mTOR inhibitors (copanlisib, idelalisib)",
    "CREBBP": "EZH2i / HDAC inhibitors",
    "BCL2": "Venetoclax",
    "JAK1": "Ruxolitinib (JAK inhibitor)",
    "STAT3": "JAK/STAT inhibitors (ruxolitinib)",
    "MTOR": "mTOR inhibitors (everolimus)",
    "ATM": "PARP inhibitors (investigational)",
    "BTK": "Ibrutinib, Zanubrutinib",
}

# Interpretation text for each case
INTERPRETATIONS = {
    "01-JJH_10679562": (
        "EZH2 p.Tyr646Phe (c.1937A>T, VAF 38.10%) 및 EZH2 p.Tyr646Asn (c.1936T>A, VAF 0.22%) "
        "gain-of-function 변이가 검출되었습니다. EZH2 Y646 변이는 germinal center B-cell (GCB) subtype "
        "DLBCL에서 특징적으로 나타나는 변이로, H3K27 trimethylation을 증가시켜 종양 발생에 기여합니다. "
        "Tazemetostat (EZH2 inhibitor)가 FDA 승인 (2020, follicular lymphoma)되어 있으며, "
        "EZH2 mutant relapsed/refractory DLBCL에서도 임상적 이점이 보고되었습니다 "
        "(Morschhauser et al., Lancet Oncol. 2020).\n\n"
        "CREBBP p.Arg1446His (VAF 39.70%)는 likely pathogenic 변이로, HAT domain의 기능 상실을 "
        "유발합니다. CREBBP 변이는 GCB-DLBCL의 30-40%에서 발견되며, MHC class II 발현 저하와 면역 회피에 관여합니다.\n\n"
        "ARID1A p.Arg1989Ter (VAF 39.41%)는 SWI/SNF chromatin remodeling complex의 기능 상실을 "
        "유발하는 pathogenic nonsense 변이입니다. DNA damage repair 결함을 초래하여 PARP inhibitor, "
        "ATR inhibitor 등의 치료 표적이 될 수 있습니다.\n\n"
        "PTEN에서 4개의 독립적인 pathogenic/likely pathogenic 변이가 검출되었습니다 "
        "(p.Phe81Cys, p.His123Arg, p.Asp52del, splice donor c.1026+1G>T). "
        "PTEN 기능 상실은 PI3K/AKT/mTOR 경로의 구성적 활성화를 유발하며, "
        "PI3K inhibitors (copanlisib, idelalisib) 또는 mTOR inhibitors (everolimus)의 치료적 표적이 됩니다.\n\n"
        "KMT2D에서 22개 이상의 변이가 검출되었으며, GNA13 p.Tyr343LeufsTer19 (VAF 29.67%), "
        "BCL2 p.Phe112Leu (VAF 29.53%), TCF3 (VAF 40.99-50.24%), MEF2B p.Leu269Val (VAF 21.19%) "
        "변이가 동시에 검출되어 GCB-DLBCL (EZB/C3 subtype)에 부합합니다 "
        "(Chapuy et al., Nat Med. 2018; Schmitz et al., NEJM. 2018)."
    ),
    "02-OKW_10673102": (
        "NOTCH2 p.Arg2400Ter (VAF 0.57%)는 pathogenic으로 분류된 truncating 변이로, "
        "NOTCH signaling의 기능 상실을 유발합니다. 추가로 NOTCH2 frameshift "
        "(p.Pro6ArgfsTer27, VAF 14.48%) 및 splice acceptor (c.74-2A>G, VAF 9.80%) "
        "변이도 검출되어 biallelic NOTCH2 inactivation이 시사됩니다.\n\n"
        "DNMT3A p.Val636Met (VAF 0.31%)는 likely pathogenic 변이로, methyltransferase domain에 "
        "위치합니다. 추가 DNMT3A 변이 (p.Cys666Tyr, VAF 1.10%)가 동시 검출되었습니다. "
        "DNMT3A 변이는 clonal hematopoiesis of indeterminate potential (CHIP)의 가장 흔한 "
        "원인 유전자이며, matched tumor tissue를 이용한 germline 확인이 권장됩니다.\n\n"
        "STAT3 p.Glu415Gln (VAF 0.05%)는 SH2 domain에 위치한 gain-of-function 변이로, "
        "JAK/STAT 경로의 구성적 활성화를 유발합니다. JAK inhibitors (ruxolitinib) 치료 표적이 됩니다.\n\n"
        "TET2에서 3개의 변이 (p.Gly1913Val, p.Tyr1255Ter, p.Gln1030Ter)가 검출되었습니다. "
        "DNMT3A + TET2 동시 변이는 clonal hematopoiesis 기원의 가능성을 높이므로 주의가 필요합니다.\n\n"
        "TCF3 p.Gly431Ser (VAF 42.50%), MEF2B p.Leu269Val (VAF 15.64%), KMT2D 변이 클러스터의 "
        "동시 검출은 GCB-DLBCL 분자 아형에 부합합니다."
    ),
    "03-OJS_23953884": (
        "Tier 1 변이가 검출되지 않았으며, 전반적으로 낮은 변이 부담(mutational burden)을 보입니다.\n\n"
        "BCL6 p.Ala517Ile (VAF 39.13%)는 germinal center reaction의 master regulator인 "
        "BCL6의 missense 변이로, DLBCL 발생에 핵심적인 역할을 합니다.\n\n"
        "MEF2B에서 2개의 변이 (p.Leu269Val VAF 16.52%, p.Arg3Ser VAF 0.24%)가 검출되었으며, "
        "p.Leu269Val는 4개 환자 모두에서 반복 검출되는 GCB-DLBCL 특이 hotspot 변이입니다.\n\n"
        "TCF3 p.Gly431Ser (VAF 2.01%), NOTCH2 p.Pro6ArgfsTer27 (VAF 18.93%), "
        "KMT2D 변이 클러스터 (5개 missense, VAF 0.29-2.48%)가 검출되었습니다.\n\n"
        "MEF2B, TCF3, KMT2D, BCL6 변이의 동시 검출은 GCB-DLBCL 분자 아형에 부합합니다. "
        "Tier 1 변이의 부재와 전반적으로 낮은 VAF 패턴은 낮은 ctDNA 방출량 또는 "
        "조기/관해 상태의 질환을 시사할 수 있으며, 강한 근거의 표적 치료 적응증은 제한적입니다."
    ),
    "04-PHJ_10680696": (
        "Tier 1 변이가 검출되지 않았으나, 다수의 Tier 2 드라이버 변이가 존재합니다.\n\n"
        "MEF2B에서 2개의 고빈도 변이 (p.Leu269Val VAF 19.84%, p.Gln279Pro VAF 13.18%)가 "
        "검출되었습니다. 두 변이의 높은 VAF는 주요 종양 클론의 드라이버 변이임을 시사합니다. "
        "MEF2B는 BCL6 전사 활성화의 핵심 조절인자로, gain-of-function 변이는 "
        "germinal center 분화를 억제합니다.\n\n"
        "CREBBP p.Gln1950Pro (VAF 10.00%)는 HAT domain에서 검출된 변이로, "
        "GCB-DLBCL의 30-40%에서 발견됩니다. CREBBP 변이 DLBCL은 EZH2 inhibitor "
        "(tazemetostat) 또는 HDAC inhibitor에 대한 반응성이 보고되었습니다.\n\n"
        "SPEN에서 3개의 frameshift 변이 (동일 codon Ser2305)가 검출되어 "
        "해당 영역이 mutational hotspot임을 시사합니다.\n\n"
        "ATM p.Arg924Gln (VAF 0.13%)는 DNA damage response 유전자 변이로 PARP inhibitor "
        "감수성과 관련될 수 있습니다. CIITA p.Ala723Ser (VAF 0.03%)는 MHC class II 전사 조절인자 "
        "변이로 면역 회피(immune evasion)에 기여합니다.\n\n"
        "MEF2B, CREBBP, KMT2D 변이의 동시 검출은 GCB-DLBCL 분자 아형에 부합하며, "
        "CIITA 변이의 동시 존재는 immune evasion phenotype을 시사합니다."
    ),
}


def clone_row(table, row_idx):
    """Deep-copy a row from a table and append it."""
    template_row = table.rows[row_idx]._tr
    new_row = copy.deepcopy(template_row)
    table._tbl.append(new_row)
    return table.rows[-1]


def set_cell_text(cell, text, font_size=Pt(9), bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    """Set cell text with formatting, clearing existing content."""
    # Clear all existing runs
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    # Always add a fresh run to avoid empty-runs issues
    run = p.add_run(str(text))
    run.font.size = font_size
    run.font.bold = bold
    run.font.name = "Malgun Gothic"


def read_tiered_csv(path):
    """Read tiered report CSV, return list of dicts."""
    rows = []
    with open(path, encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for r in reader:
            rows.append(r)
    return rows


def read_annotated_csv(path):
    """Read annotated CSV for tier counts."""
    tier_counts = {"Tier 1": 0, "Tier 2": 0, "Tier 3": 0, "Tier 4": 0}
    total = 0
    with open(path, encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for r in reader:
            total += 1
            tier = r.get("tier", "Tier 4")
            if tier in tier_counts:
                tier_counts[tier] += 1
    return total, tier_counts


def generate_report(case_id, tiered_csv, annotated_csv, output_path):
    """Generate a .docx report for one case."""
    meta = CASES[case_id]
    tiered = read_tiered_csv(tiered_csv)
    total, tier_counts = read_annotated_csv(annotated_csv)

    tier12 = [r for r in tiered if r["ASCO/AMP Classification"] in ("Tier 1", "Tier 2")]
    tier3 = [r for r in tiered if r["ASCO/AMP Classification"] == "Tier 3"]

    doc = Document(TEMPLATE)

    # --- Table 0: Patient info ---
    t0 = doc.tables[0]
    sz = Pt(9)
    # Row 0: patient/registration, specimen type, test number
    set_cell_text(t0.rows[0].cells[1], meta["patient"], sz)
    set_cell_text(t0.rows[0].cells[3], meta["specimen"], sz)
    set_cell_text(t0.rows[0].cells[5], meta["reg_no"], sz)
    # Row 1: reg number, specimen date, test date
    set_cell_text(t0.rows[1].cells[1], meta["reg_no"], sz)
    set_cell_text(t0.rows[1].cells[3], "", sz)
    set_cell_text(t0.rows[1].cells[5], meta["test_date"], sz)
    # Row 2: physician, specimen received, report date
    set_cell_text(t0.rows[2].cells[1], "", sz)
    set_cell_text(t0.rows[2].cells[3], "", sz)
    set_cell_text(t0.rows[2].cells[5], meta["test_date"], sz)
    # Row 3: requesting physician, specimen type detail, final report date
    set_cell_text(t0.rows[3].cells[1], "", sz)
    set_cell_text(t0.rows[3].cells[3], meta["specimen"], sz)
    set_cell_text(t0.rows[3].cells[5], meta["test_date"], sz)

    # --- Table 3: Tier 1/2 variants ---
    t3 = doc.tables[3]
    # Remove existing data rows (rows 2..end), keep header rows 0,1
    while len(t3.rows) > 2:
        t3._tbl.remove(t3.rows[2]._tr)

    # Add tier 1/2 data rows
    for v in tier12:
        new_row = clone_row(t3, 1)  # clone from header row for structure
        gene = v["Gene"]
        therapeutic = THERAPEUTIC.get(gene, "")
        values = [
            v["ASCO/AMP Classification"],
            gene,
            v.get("Canonical name", ""),
            v.get("Accession", ""),
            v.get("Nucleotide change", ""),
            v.get("AA change", ""),
            v.get("% VAF", ""),
            therapeutic,
        ]
        for i, val in enumerate(values):
            set_cell_text(new_row.cells[i], val, font_size=Pt(9), bold=False)

    # Remove the cloned header-style row if no data
    if not tier12:
        new_row = clone_row(t3, 1)
        for i in range(8):
            set_cell_text(new_row.cells[i], "", font_size=Pt(9))

    # --- Table 4: Tier 3 (VUS) ---
    t4 = doc.tables[4]
    while len(t4.rows) > 2:
        t4._tbl.remove(t4.rows[2]._tr)

    for v in tier3:
        new_row = clone_row(t4, 1)
        values = [
            v["ASCO/AMP Classification"],
            v["Gene"],
            v.get("Canonical name", ""),
            v.get("Accession", ""),
            v.get("Nucleotide change", ""),
            v.get("AA change", ""),
            v.get("% VAF", ""),
            "",
        ]
        for i, val in enumerate(values):
            set_cell_text(new_row.cells[i], val, font_size=Pt(9), bold=False)

    # --- Table 6: Interpretation ---
    t6 = doc.tables[6]
    cell = t6.rows[0].cells[0]
    # Clear all existing paragraphs content
    for p in cell.paragraphs:
        p_el = p._element
        for child in list(p_el):
            if child.tag.endswith('}r'):
                p_el.remove(child)
    interp = INTERPRETATIONS.get(case_id, "")
    p = cell.paragraphs[0]
    run = p.add_run(interp)
    run.font.size = Pt(9)
    run.font.name = "Malgun Gothic"

    # --- Table 8: Limitations ---
    t8 = doc.tables[8]
    cell8 = t8.rows[0].cells[0]
    for p in cell8.paragraphs:
        p_el = p._element
        for child in list(p_el):
            if child.tag.endswith('}r'):
                p_el.remove(child)
    lim_text = (
        "본 검사는 차세대염기서열분석법으로 시행되었으며 SNP, Small indel 의 검출이 가능하고, "
        "CNV 및 gene rearrangement 등의 구조 변이는 검출이 불가능합니다.\n"
        "본 검사는 Germline 변이와 Somatic 변이를 구분할 수 없으며 "
        "Variant allele frequency 50%, 100% 에 가까운 경우 Germline variant 의 가능성을 고려해야 합니다.\n"
        "검출 변이의 임상적 중요성에 따라 Tier I, II, III에 해당하는 변이만 보고하였습니다. "
        "(J Mol Diagn. 2017;19;4-23)\n"
        "* Tier I : 임상적 의미가 확인된 변이, "
        "** Tier II : 잠재적 임상적 의미가 있는 변이, "
        "*** Tier III : 임상적 의미가 알려지지 않은 변이\n\n"
        f"Total variants analyzed: {total} | "
        f"Tier 1: {tier_counts['Tier 1']} | Tier 2: {tier_counts['Tier 2']} | "
        f"Tier 3: {tier_counts['Tier 3']} | Tier 4: {tier_counts['Tier 4']}"
    )
    p8 = cell8.paragraphs[0]
    run8 = p8.add_run(lim_text)
    run8.font.size = Pt(8)
    run8.font.name = "Malgun Gothic"

    doc.save(output_path)
    print(f"  Generated: {output_path}")
    print(f"    Tier 1/2: {len(tier12)} variants, Tier 3: {len(tier3)} variants")


def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    base_dir = os.path.dirname(script_dir)  # ctDNA root
    data_dir = os.path.join(base_dir, "0224")

    for case_id in CASES:
        case_dir = os.path.join(data_dir, case_id)
        tiered_csv = os.path.join(case_dir, f"{case_id}_tiered_report.csv")
        annotated_csv = os.path.join(case_dir, f"{case_id}_annotated.csv")
        output_path = os.path.join(case_dir, f"{case_id}_clinical_report.docx")

        if not os.path.exists(tiered_csv):
            print(f"  SKIP {case_id}: {tiered_csv} not found")
            continue
        if not os.path.exists(annotated_csv):
            print(f"  SKIP {case_id}: {annotated_csv} not found")
            continue

        print(f"Processing {case_id}...")
        generate_report(case_id, tiered_csv, annotated_csv, output_path)

    print("\nDone!")


if __name__ == "__main__":
    main()
