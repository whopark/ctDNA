#!/usr/bin/env python3
"""
Generate clinical .docx reports from tiered CSV + annotated CSV data
for the 0325 batch using the provided template.docx.
"""

import csv
import copy
import sys
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

TEMPLATE = os.path.join(os.path.dirname(__file__), 'template.docx')

# Case metadata - patient identifiers anonymized
CASES = {
    "01-": {
        "patient": "Patient 01",
        "reg_no": "0325-01",
        "specimen": "Peripheral Blood",
        "test_date": "2026-03-25",
    },
    "02-": {
        "patient": "Patient 02",
        "reg_no": "0325-02",
        "specimen": "Peripheral Blood",
        "test_date": "2026-03-25",
    },
    "03-": {
        "patient": "Patient 03",
        "reg_no": "0325-03",
        "specimen": "Peripheral Blood",
        "test_date": "2026-03-25",
    },
    "04-": {
        "patient": "Patient 04",
        "reg_no": "0325-04",
        "specimen": "Peripheral Blood",
        "test_date": "2026-03-25",
    },
    "05-": {
        "patient": "Patient 05",
        "reg_no": "0325-05",
        "specimen": "Peripheral Blood",
        "test_date": "2026-03-25",
    },
}

# Therapeutic implications lookup
THERAPEUTIC = {
    "EZH2": "Tazemetostat (FDA-approved, FL; R/R DLBCL data)",
    "PTEN": "PI3K/AKT/mTOR inhibitors (copanlisib, idelalisib)",
    "CREBBP": "EZH2i / HDAC inhibitors",
    "BCL2": "Venetoclax",
    "JAK1": "Ruxolitinib (JAK inhibitor)",
    "STAT3": "JAK/STAT inhibitors (ruxolitinib)",
    "MTOR": "mTOR inhibitors (everolimus)",
    "ATM": "PARP inhibitors (investigational)",
    "BTK": "Ibrutinib, Zanubrutinib",
    "MYD88": "BTK inhibitors (ibrutinib) - MYD88 L265P/L273P",
    "CD79B": "BTK inhibitors (ibrutinib)",
    "TP53": "DNA damage agents; consider venetoclax combinations",
    "MYC": "Dose-adjusted regimens (DA-EPOCH-R)",
    "BRAF": "BRAF inhibitors (vemurafenib/dabrafenib) - context dependent",
}

# Interpretation text for each case
INTERPRETATIONS = {
    "01-": (
        "MYD88 p.Leu273Pro (c.818T>C, VAF 12.93%) pathogenic 변이가 검출되었습니다. "
        "MYD88 L265P/L273P 변이는 ABC-DLBCL (activated B-cell type)의 가장 특징적인 변이로, "
        "약 30%의 ABC-DLBCL에서 발견됩니다. MYD88 변이 양성 DLBCL은 BTK inhibitor (ibrutinib)에 대한 "
        "반응률이 높은 것으로 보고됩니다 (Wilson et al., Nat Med. 2015).\n\n"
        "CD79B p.Tyr197Ser (VAF 19.55%) 변이가 동시 검출되었습니다. MYD88 + CD79B 동시 변이는 "
        "MCD (MYD88/CD79B double-mutant) molecular subtype을 정의하며, BTK inhibitor에 대한 "
        "반응률이 특히 높습니다 (약 80% ORR, Schmitz et al., NEJM 2018).\n\n"
        "PIM1에서 다수의 missense 변이 (9개)가 검출되었으며, 이는 aberrant somatic hypermutation (aSHM)의 "
        "결과로 해석됩니다. PIM1 hypermutation은 ABC-DLBCL에서 흔히 관찰됩니다.\n\n"
        "NOTCH2 frameshift (p.Pro6ArgfsTer27, VAF 13.68%), CARD11 p.Ala630Pro (VAF 10.35%), "
        "KMT2D 변이 클러스터, B2M frameshift (p.Thr24LeufsTer27, VAF 7.56%)가 검출되었습니다. "
        "B2M 변이는 MHC class I 발현 저하로 인한 면역 회피 기전에 관여합니다.\n\n"
        "종합적으로 MYD88 + CD79B 동시 변이, PIM1 hypermutation, NOTCH2/CARD11 변이 패턴은 "
        "ABC-DLBCL (MCD subtype)에 부합하며, BTK inhibitor 기반 치료의 적응증이 됩니다."
    ),
    "02-": (
        "본 검체에서 Tier 1 변이는 검출되지 않았으나, 다수의 Tier 2 드라이버 변이가 존재합니다.\n\n"
        "PIM1에서 11개의 missense 및 nonsense 변이가 검출되었으며, 높은 VAF (최대 86.04%)를 보입니다. "
        "이는 aberrant somatic hypermutation (aSHM)의 전형적인 패턴으로, DLBCL의 발생 기전에 관여합니다.\n\n"
        "GNA13에서 3개의 변이 (splice site, p.Lys61Arg, p.Gln27His)가 검출되었습니다 (VAF ~40%). "
        "GNA13 변이는 GCB-DLBCL의 약 25%에서 발견되며, germinal center confinement 결함과 관련됩니다.\n\n"
        "MYC p.Arg80Gly (VAF 33.37%) 변이가 검출되었습니다. MYC transactivation domain 변이는 "
        "단백질 안정성을 증가시켜 종양 발생에 기여합니다.\n\n"
        "CD79B frameshift (p.Ile175LeufsTer35, VAF 32.74%), CD58 frameshift (p.Lys184AsnfsTer6, VAF 30.66%)가 "
        "검출되었습니다. CD58 기능 상실은 NK cell 매개 면역 회피 기전의 핵심입니다.\n\n"
        "MEF2B 3개 변이 (p.Met1?, p.Leu269Val, p.Gln279Pro), CREBBP, ARID1A, BRAF p.Leu711Ile 변이가 "
        "동시 검출되었습니다.\n\n"
        "GNA13 + MEF2B + PIM1 hypermutation 패턴은 GCB-DLBCL에 부합하며, "
        "MYC 변이와 CD58 면역 회피 기전이 동시에 존재합니다."
    ),
    "03-": (
        "MYD88 p.Leu273Pro (VAF 58.29%)가 주요 드라이버 변이로 검출되었습니다. 높은 VAF는 "
        "주요 종양 클론의 드라이버 변이임을 시사합니다.\n\n"
        "TP53 p.Ile232Asn (VAF 1.35%)이 검출되었습니다. TP53 변이는 DLBCL의 예후 불량 인자로, "
        "표준 R-CHOP 치료에 대한 반응률이 낮습니다.\n\n"
        "ATM에서 2개의 truncating 변이 (p.Ser2192Ter VAF 7.27%, p.Ser421Ter VAF 3.89%)가 검출되어 "
        "biallelic ATM inactivation이 시사됩니다. ATM 기능 상실은 DNA damage response 결함을 유발하며, "
        "PARP inhibitor 감수성과 관련될 수 있습니다.\n\n"
        "RB1 p.Glu746Ter (VAF 10.16%)는 tumor suppressor 기능 상실 변이입니다. "
        "추가로 RB1 missense 변이가 다수 검출되어 RB1 pathway 비활성화가 확인됩니다.\n\n"
        "CREBBP에서 2개의 splice/missense 변이 (c.4395-2A>T, p.Gln1491Lys)가 검출되었으며, "
        "KMT2D frameshift (p.Ala2119ArgfsTer36, VAF 8.50%)도 확인됩니다.\n\n"
        "TET2 p.Gln341Ter (VAF 92.99%)는 매우 높은 VAF를 보이며, clonal hematopoiesis of "
        "indeterminate potential (CHIP) 기원 변이의 가능성을 고려해야 합니다.\n\n"
        "MYD88 + TP53 + ATM + RB1 동시 변이는 고위험 ABC-DLBCL을 시사하며, "
        "BTK inhibitor 기반 치료와 DNA damage targeting 전략을 고려할 수 있습니다."
    ),
    "04-": (
        "ATM p.Glu2039Lys (VAF 3.06%)가 Tier 1 pathogenic 변이로 검출되었습니다. "
        "ATM kinase domain 변이는 DNA damage response 결함을 유발합니다.\n\n"
        "TCF3 p.Gly431Ser (VAF 40.23%)는 E2A/TCF3 전사인자의 DNA binding domain 변이로, "
        "GCB-DLBCL에서 흔히 발견됩니다.\n\n"
        "MEF2B p.Leu269Val (VAF 14.56%)는 4개 환자 모두에서 반복 검출되는 GCB-DLBCL 특이 "
        "hotspot 변이입니다. MEF2B는 BCL6 전사 활성화의 핵심 조절인자입니다.\n\n"
        "NOTCH2 frameshift (p.Pro6ArgfsTer27, VAF 14.45%), CREBBP p.Val1956Gly (VAF 10.53%), "
        "TET2 p.Gln964Ter (VAF 5.74%)가 검출되었습니다. TET2 truncating 변이는 "
        "CHIP 기원 가능성을 고려해야 합니다.\n\n"
        "STAT6 p.Gln65Pro (VAF 4.89%)는 JAK/STAT pathway 활성화 변이입니다.\n\n"
        "KMT2C/KMT2D 변이 클러스터가 검출되어 히스톤 methylation 조절 이상이 확인됩니다.\n\n"
        "TCF3, MEF2B, KMT2D 변이의 동시 검출은 GCB-DLBCL 분자 아형에 부합합니다. "
        "ATM 변이는 PARP inhibitor 감수성과 관련될 수 있습니다."
    ),
    "05-": (
        "TP53 p.Leu145Gln (VAF 43.25%)이 주요 드라이버 변이로 검출되었습니다. DNA binding domain "
        "변이는 tumor suppressor 기능 상실을 유발하며, DLBCL의 독립적 예후 불량 인자입니다.\n\n"
        "MYC에서 다수의 변이가 검출되었습니다: p.Glu54Asp (VAF 35.19%, Tier 1), "
        "p.Tyr89Asp (VAF 40.56%), p.Glu44Asp (VAF 37.11%), p.Gln49His (VAF 34.84%) 등. "
        "MYC N-terminal transactivation domain 변이 클러스터는 단백질 안정성 증가와 "
        "oncogenic activity 강화를 유발합니다.\n\n"
        "ID3에서 splice site (c.300+1G>C, VAF 41.34%) 및 p.Leu64Phe (VAF 40.95%) 변이가 검출되었습니다. "
        "ID3 기능 상실 변이는 Burkitt lymphoma의 70%에서 발견되며, TCF3 활성을 증가시켜 "
        "MYC 발현을 촉진합니다.\n\n"
        "RB1 p.Glu746Ter (VAF 10.84%) truncating 변이와 DDX3X에서 3개의 pathogenic 변이 "
        "(p.Gln14Ter, p.Arg488His, p.Pro274Ser)가 검출되었습니다. DDX3X는 RNA helicase로 "
        "번역 조절에 관여합니다.\n\n"
        "GNA13 p.Val92Glu (VAF 60.80%), CCND3 변이 클러스터 (p.Val67Gly, splice, p.Cys73Gly)가 "
        "동시 검출되었습니다. CCND3 C-terminal 변이는 단백질 안정성 증가를 유발합니다.\n\n"
        "TP53 + MYC + ID3 + CCND3 동시 변이 패턴은 high-grade B-cell lymphoma (HGBL) 또는 "
        "Burkitt-like lymphoma를 강력히 시사합니다. MYC rearrangement에 대한 FISH 검사가 권장되며, "
        "고강도 항암요법 (DA-EPOCH-R 또는 CODOX-M/IVAC)을 고려해야 합니다."
    ),
}


def clone_row(table, row_idx):
    """Deep-copy a row from a table and append it."""
    template_row = table.rows[row_idx]._tr
    new_row = copy.deepcopy(template_row)
    table._tbl.append(new_row)
    return table.rows[-1]


def set_cell_text(cell, text, font_size=Pt(9), bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    """Set cell text with formatting, clearing existing content."""
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.size = font_size
    run.font.bold = bold
    run.font.name = "Malgun Gothic"


def read_tiered_csv(path):
    """Read tiered report CSV, return list of dicts."""
    rows = []
    with open(path, encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for r in reader:
            rows.append(r)
    return rows


def read_annotated_csv(path):
    """Read annotated CSV for tier counts."""
    tier_counts = {"Tier 1": 0, "Tier 2": 0, "Tier 3": 0, "Tier 4": 0}
    total = 0
    with open(path, encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for r in reader:
            total += 1
            tier = r.get("tier", "Tier 4")
            if tier in tier_counts:
                tier_counts[tier] += 1
    return total, tier_counts


def generate_report(case_id, tiered_csv, annotated_csv, output_path):
    """Generate a .docx report for one case."""
    meta = CASES[case_id]
    tiered = read_tiered_csv(tiered_csv)
    total, tier_counts = read_annotated_csv(annotated_csv)

    tier1 = [r for r in tiered if r["ASCO/AMP Classification"] == "Tier 1"]
    tier2 = [r for r in tiered if r["ASCO/AMP Classification"] == "Tier 2"]
    tier12 = tier1 + tier2
    tier3 = [r for r in tiered if r["ASCO/AMP Classification"] == "Tier 3"]

    doc = Document(TEMPLATE)

    # --- Table 0: Patient info ---
    t0 = doc.tables[0]
    sz = Pt(9)
    set_cell_text(t0.rows[0].cells[1], meta["patient"], sz)
    set_cell_text(t0.rows[0].cells[3], meta["specimen"], sz)
    set_cell_text(t0.rows[0].cells[5], meta["reg_no"], sz)
    set_cell_text(t0.rows[1].cells[1], meta["reg_no"], sz)
    set_cell_text(t0.rows[1].cells[3], "", sz)
    set_cell_text(t0.rows[1].cells[5], meta["test_date"], sz)
    set_cell_text(t0.rows[2].cells[1], "", sz)
    set_cell_text(t0.rows[2].cells[3], "", sz)
    set_cell_text(t0.rows[2].cells[5], meta["test_date"], sz)
    set_cell_text(t0.rows[3].cells[1], "", sz)
    set_cell_text(t0.rows[3].cells[3], meta["specimen"], sz)
    set_cell_text(t0.rows[3].cells[5], meta["test_date"], sz)

    # --- Table 3: Tier 1/2 variants ---
    t3 = doc.tables[3]
    while len(t3.rows) > 2:
        t3._tbl.remove(t3.rows[2]._tr)

    for v in tier12:
        new_row = clone_row(t3, 1)
        gene = v["Gene"]
        therapeutic = THERAPEUTIC.get(gene, "")
        values = [
            v["ASCO/AMP Classification"],
            gene,
            v.get("Canonical name", ""),
            v.get("Accession", ""),
            v.get("Nucleotide change", ""),
            v.get("AA change", ""),
            v.get("% VAF", ""),
            therapeutic,
        ]
        for i, val in enumerate(values):
            set_cell_text(new_row.cells[i], val, font_size=Pt(9), bold=False)

    if not tier12:
        new_row = clone_row(t3, 1)
        for i in range(8):
            set_cell_text(new_row.cells[i], "No Tier 1/2 variants detected", font_size=Pt(9))

    # --- Table 4: Tier 3 (VUS) ---
    t4 = doc.tables[4]
    while len(t4.rows) > 2:
        t4._tbl.remove(t4.rows[2]._tr)

    # Limit Tier 3 to top 20 variants (sorted by VAF descending) to keep report manageable
    tier3_sorted = sorted(tier3, key=lambda x: float(x.get("% VAF", "0%").rstrip("%")), reverse=True)[:20]

    for v in tier3_sorted:
        new_row = clone_row(t4, 1)
        values = [
            v["ASCO/AMP Classification"],
            v["Gene"],
            v.get("Canonical name", ""),
            v.get("Accession", ""),
            v.get("Nucleotide change", ""),
            v.get("AA change", ""),
            v.get("% VAF", ""),
            "",
        ]
        for i, val in enumerate(values):
            set_cell_text(new_row.cells[i], val, font_size=Pt(9), bold=False)

    # --- Table 6: Interpretation ---
    t6 = doc.tables[6]
    cell = t6.rows[0].cells[0]
    for p in cell.paragraphs:
        p_el = p._element
        for child in list(p_el):
            if child.tag.endswith('}r'):
                p_el.remove(child)
    interp = INTERPRETATIONS.get(case_id, "")
    p = cell.paragraphs[0]
    run = p.add_run(interp)
    run.font.size = Pt(9)
    run.font.name = "Malgun Gothic"

    # --- Table 8: Limitations ---
    t8 = doc.tables[8]
    cell8 = t8.rows[0].cells[0]
    for p in cell8.paragraphs:
        p_el = p._element
        for child in list(p_el):
            if child.tag.endswith('}r'):
                p_el.remove(child)
    lim_text = (
        "본 검사는 차세대염기서열분석법으로 시행되었으며 SNP, Small indel 의 검출이 가능하고, "
        "CNV 및 gene rearrangement 등의 구조 변이는 검출이 불가능합니다.\n"
        "본 검사는 Germline 변이와 Somatic 변이를 구분할 수 없으며 "
        "Variant allele frequency 50%, 100% 에 가까운 경우 Germline variant 의 가능성을 고려해야 합니다.\n"
        "검출 변이의 임상적 중요성에 따라 Tier I, II, III에 해당하는 변이만 보고하였습니다. "
        "(J Mol Diagn. 2017;19;4-23)\n"
        "* Tier I : 임상적 의미가 확인된 변이, "
        "** Tier II : 잠재적 임상적 의미가 있는 변이, "
        "*** Tier III : 임상적 의미가 알려지지 않은 변이\n\n"
        f"Total variants analyzed: {total} | "
        f"Tier 1: {tier_counts['Tier 1']} | Tier 2: {tier_counts['Tier 2']} | "
        f"Tier 3: {tier_counts['Tier 3']} | Tier 4: {tier_counts['Tier 4']}"
    )
    p8 = cell8.paragraphs[0]
    run8 = p8.add_run(lim_text)
    run8.font.size = Pt(8)
    run8.font.name = "Malgun Gothic"

    doc.save(output_path)
    print(f"  Generated: {output_path}")
    print(f"    Tier 1: {len(tier1)} | Tier 2: {len(tier2)} | Tier 3: {len(tier3)} variants")


def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))

    for case_id in CASES:
        case_dir = os.path.join(script_dir, case_id)
        tiered_csv = os.path.join(case_dir, f"{case_id}_tiered_report.csv")
        annotated_csv = os.path.join(case_dir, f"{case_id}_annotated.csv")
        output_path = os.path.join(script_dir, f"{case_id}_clinical_report.docx")

        if not os.path.exists(tiered_csv):
            print(f"  SKIP {case_id}: {tiered_csv} not found")
            continue
        if not os.path.exists(annotated_csv):
            print(f"  SKIP {case_id}: {annotated_csv} not found")
            continue

        print(f"Processing {case_id}...")
        generate_report(case_id, tiered_csv, annotated_csv, output_path)

    print("\nDone! Reports generated in 0325/ directory.")


if __name__ == "__main__":
    main()
