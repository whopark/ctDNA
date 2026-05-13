#!/usr/bin/env python3
"""Generate clinical .docx reports per SPEC-REPORT-001.

Patient meta is loaded from meta.json (case_meta.py), interpretation
text from interpretations.yaml (interpretations_loader.py), and QC
stats from the annotated CSV dp column (qc_stats.py).
"""

import csv
import copy
import os
import sys
import site
from pathlib import Path


def _add_user_site_packages():
    """Add user site-packages paths so python-docx resolves without a venv."""
    try:
        user_site = site.getusersitepackages()
        if user_site and user_site not in sys.path:
            sys.path.append(user_site)
    except Exception:
        pass
    home = Path.home()
    base = home / "AppData/Local/Packages"
    for ver in ("3.13", "3.12", "3.11"):
        suffix = ver.replace(".", "")
        sp = base / f"PythonSoftwareFoundation.Python.{ver}_qbz5n2kfra8p0/LocalCache/local-packages/Python{suffix}/site-packages"
        if sp.exists() and str(sp) not in sys.path:
            sys.path.append(str(sp))


_add_user_site_packages()

from docx import Document
from docx.shared import Pt

_REPO_ROOT = os.path.abspath(os.path.dirname(__file__))
if _REPO_ROOT not in sys.path:
    sys.path.insert(0, _REPO_ROOT)

from case_meta import load_case_meta  # noqa: E402
from qc_stats import compute_qc  # noqa: E402
from interpretations_loader import load_interpretations, resolve_interpretation  # noqa: E402
from report_tables import (  # noqa: E402
    set_cell_text, fill_table0_patient_info,
    fill_table14_qc_data, fill_table15_signatures, vaf_float,
)

# Self-updating knowledge base (optional dependency).
try:
    from kb import load_kb, therapeutic_map, interpretations_map
    _KB = load_kb()
    _KB_THERAPEUTICS = therapeutic_map(_KB)
    _KB_INTERPRETATIONS = interpretations_map(_KB)
except Exception:  # pragma: no cover — KB is optional
    _KB = None
    _KB_THERAPEUTICS = {}
    _KB_INTERPRETATIONS = {}


def _find_template(template_path: str | None = None) -> str:
    """Locate template.docx.

    Priority order (REQ-9):
      1. explicit template_path argument (if it exists)
      2. ./template.docx next to this script
      3. ./0325/template.docx
      4. parent-dir template.docx (legacy fallback)
    """
    if template_path and os.path.exists(template_path):
        return template_path
    here = os.path.dirname(os.path.abspath(__file__))
    candidates = [
        os.path.join(here, "template.docx"),
        os.path.join(here, "0325", "template.docx"),
        os.path.join(os.path.dirname(here), "template.docx"),
    ]
    for p in candidates:
        if os.path.exists(p):
            return p
    # Last resort: let python-docx raise a clear error
    return candidates[0]


TEMPLATE = _find_template()

# Therapeutic implications lookup (fallback when KB has no entry for the gene)
THERAPEUTIC = {
    "EZH2": "Tazemetostat (FDA-approved, FL; R/R DLBCL data)",
    "PTEN": "PI3K/AKT/mTOR inhibitors (copanlisib, idelalisib)",
    "CREBBP": "EZH2i / HDAC inhibitors", "BCL2": "Venetoclax",
    "JAK1": "Ruxolitinib (JAK inhibitor)",
    "STAT3": "JAK/STAT inhibitors (ruxolitinib)",
    "MTOR": "mTOR inhibitors (everolimus)",
    "ATM": "PARP inhibitors (investigational)",
    "BTK": "Ibrutinib, Zanubrutinib",
    "MYD88": "BTK inhibitors (ibrutinib) - MYD88 L265P/L273P",
    "CD79B": "BTK inhibitors (ibrutinib)",
    "TP53": "DNA damage agents; consider venetoclax combinations",
    "MYC": "Dose-adjusted regimens (DA-EPOCH-R)",
    "BRAF": "BRAF inhibitors (vemurafenib/dabrafenib) - context dependent",
}


def clone_row(table, row_idx):
    """Deep-copy a row from a table and append it."""
    template_row = table.rows[row_idx]._tr
    new_row = copy.deepcopy(template_row)
    table._tbl.append(new_row)
    return table.rows[-1]


def read_tiered_csv(path):
    """Read tiered report CSV, return list of dicts."""
    rows = []
    with open(path, encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for r in reader:
            rows.append(r)
    return rows


def _fill_tier12(doc, tiered):
    """Populate Table 3 with Tier 1/2 variants (REQ-7: Tier 1 first, VAF desc)."""
    tier1 = [r for r in tiered if r["ASCO/AMP Classification"] == "Tier 1"]
    tier2 = [r for r in tiered if r["ASCO/AMP Classification"] == "Tier 2"]
    tier12 = (
        sorted(tier1, key=vaf_float, reverse=True)
        + sorted(tier2, key=vaf_float, reverse=True)
    )
    t3 = doc.tables[3]
    while len(t3.rows) > 2:
        t3._tbl.remove(t3.rows[2]._tr)
    for v in tier12:
        new_row = clone_row(t3, 1)
        gene = v["Gene"]
        therapeutic = _KB_THERAPEUTICS.get(gene) or THERAPEUTIC.get(gene, "")
        values = [
            v["ASCO/AMP Classification"], gene,
            v.get("Canonical name", ""), v.get("Accession", ""),
            v.get("Nucleotide change", ""), v.get("AA change", ""),
            v.get("% VAF", ""), therapeutic,
        ]
        for i, val in enumerate(values):
            set_cell_text(new_row.cells[i], val, font_size=Pt(9), bold=False)
    if not tier12:
        new_row = clone_row(t3, 1)
        for i in range(8):
            set_cell_text(
                new_row.cells[i],
                "No Tier 1/2 variants detected",
                font_size=Pt(9),
            )
    return len(tier1), len(tier2)


def _fill_tier3(doc, tiered):
    """Populate Table 4 with up to 14 Tier 3 rows sorted by VAF desc (REQ-8)."""
    tier3 = [r for r in tiered if r["ASCO/AMP Classification"] == "Tier 3"]
    t4 = doc.tables[4]
    while len(t4.rows) > 2:
        t4._tbl.remove(t4.rows[2]._tr)
    tier3_sorted = sorted(tier3, key=vaf_float, reverse=True)[:14]
    for v in tier3_sorted:
        new_row = clone_row(t4, 1)
        values = [
            v["ASCO/AMP Classification"], v["Gene"],
            v.get("Canonical name", ""), v.get("Accession", ""),
            v.get("Nucleotide change", ""), v.get("AA change", ""),
            v.get("% VAF", ""), "",
        ]
        for i, val in enumerate(values):
            set_cell_text(new_row.cells[i], val, font_size=Pt(9), bold=False)
    return len(tier3)


def _fill_interpretation(doc, case_id):
    """Populate Table 6 interpretation text using the priority resolver (REQ-5)."""
    t6 = doc.tables[6]
    cell = t6.rows[0].cells[0]
    for p in cell.paragraphs:
        p_el = p._element
        for child in list(p_el):
            if child.tag.endswith('}r'):
                p_el.remove(child)
    yaml_path = os.path.join(_REPO_ROOT, "interpretations.yaml")
    yaml_map = load_interpretations(yaml_path)
    interp = resolve_interpretation(
        case_id, yaml_map=yaml_map, kb_map=_KB_INTERPRETATIONS,
    )
    p = cell.paragraphs[0]
    run = p.add_run(interp)
    run.font.size = Pt(9)
    run.font.name = "Malgun Gothic"


def _fill_limitations(doc, qc):
    """Populate Table 8 with limitation footer + variant counts."""
    t8 = doc.tables[8]
    cell8 = t8.rows[0].cells[0]
    for p in cell8.paragraphs:
        p_el = p._element
        for child in list(p_el):
            if child.tag.endswith('}r'):
                p_el.remove(child)
    tc = qc["tier_counts"]
    lim_text = (
        "본 검사는 차세대염기서열분석법으로 시행되었으며 SNP, Small indel 의 검출이 가능하고, "
        "CNV 및 gene rearrangement 등의 구조 변이는 검출이 불가능합니다.\n"
        "본 검사는 Germline 변이와 Somatic 변이를 구분할 수 없으며 "
        "Variant allele frequency 50%, 100% 에 가까운 경우 Germline variant 의 가능성을 고려해야 합니다.\n"
        "검출 변이의 임상적 중요성에 따라 Tier I, II, III에 해당하는 변이만 보고하였습니다. "
        "(J Mol Diagn. 2017;19;4-23)\n"
        "* Tier I : 임상적 의미가 확인된 변이, "
        "** Tier II : 잠재적 임상적 의미가 있는 변이, "
        "*** Tier III : 임상적 의미가 알려지지 않은 변이\n\n"
        f"Total variants analyzed: {qc['total_variants']} | "
        f"Tier 1: {tc['Tier 1']} | Tier 2: {tc['Tier 2']} | "
        f"Tier 3: {tc['Tier 3']} | Tier 4: {tc['Tier 4']}"
    )
    p8 = cell8.paragraphs[0]
    run8 = p8.add_run(lim_text)
    run8.font.size = Pt(8)
    run8.font.name = "Malgun Gothic"


# @AX:NOTE [AUTO] case_dir is derived from os.path.dirname(annotated_csv) so the existing
# @AX:NOTE [AUTO] 4-arg GUI callsite stays compatible after the meta.json refactor (SPEC-REPORT-001).
def generate_report(case_id, tiered_csv, annotated_csv, output_path,
                    template_path=None):
    """Generate a .docx report for one case.

    The directory containing the input CSVs is also expected to contain
    meta.json. If meta.json is missing/malformed the report still renders
    but the corresponding Table 0 / Table 15 cells are left blank (REQ-6).
    """
    case_dir = os.path.dirname(annotated_csv) or "."
    meta, warnings = load_case_meta(case_dir)
    for w in warnings:
        print(f"  [WARNING] {case_id}: {w}", file=sys.stderr)

    tiered = read_tiered_csv(tiered_csv)
    qc = compute_qc(annotated_csv)

    doc = Document(template_path or TEMPLATE)

    # Table 0: 12-cell patient meta (REQ-2, REQ-2b)
    fill_table0_patient_info(doc.tables[0], meta)

    # Tables 3/4: variant rows with ordering (REQ-7, REQ-8)
    n_tier1, n_tier2 = _fill_tier12(doc, tiered)
    n_tier3 = _fill_tier3(doc, tiered)

    # Table 6: clinical interpretation (REQ-5)
    _fill_interpretation(doc, case_id)

    # Table 8: limitations + variant counts
    _fill_limitations(doc, qc)

    # Table 14: QC stats (REQ-1)
    fill_table14_qc_data(doc.tables[14], qc)

    # Table 15: examiner / reporter signatures (REQ-3)
    fill_table15_signatures(doc.tables[15], meta)

    doc.save(output_path)
    print(f"  Generated: {output_path}")
    print(f"    Tier 1: {n_tier1} | Tier 2: {n_tier2} | Tier 3: {n_tier3} variants")


def main():
    """Iterate case folders under the script directory and render each report."""
    script_dir = os.path.dirname(os.path.abspath(__file__))
    case_prefixes = ["01-", "02-", "03-", "04-", "05-"]

    for case_id in case_prefixes:
        case_dir = os.path.join(script_dir, case_id)
        tiered_csv = os.path.join(case_dir, f"{case_id}_tiered_report.csv")
        annotated_csv = os.path.join(case_dir, f"{case_id}_annotated.csv")
        output_path = os.path.join(script_dir, f"{case_id}_clinical_report.docx")

        if not os.path.exists(tiered_csv):
            print(f"  SKIP {case_id}: {tiered_csv} not found")
            continue
        if not os.path.exists(annotated_csv):
            print(f"  SKIP {case_id}: {annotated_csv} not found")
            continue

        print(f"Processing {case_id}...")
        generate_report(case_id, tiered_csv, annotated_csv, output_path)

    print("\nDone! Reports generated.")


if __name__ == "__main__":
    main()
